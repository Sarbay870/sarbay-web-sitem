import os
import shutil
from PIL import Image

# --- DÖNÜŞTÜRME KÜTÜPHANELERİ ---
# (O 'pip install ...' komutuyla kurduklarımız)
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase import pdfmetrics
    REPORTLAB_OK = True
except ImportError: REPORTLAB_OK = False

try:
    from docx import Document
    DOCX_OK = True
except ImportError: DOCX_OK = False

try:
    from pypdf import PdfReader
    PYPDF_OK = True
except ImportError: PYPDF_OK = False

try:
    import pandas as pd
    import openpyxl
    PANDAS_OK = True
except ImportError: PANDAS_OK = False
    
try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_OK = True
except ImportError: OCR_OK = False
# --- BİTTİ ---


# --- DÖNÜŞTÜRME YARDIMCILARI (Senin .pyw kodundan uyarlandı) ---

def safe_out_path(p, new_ext):
    """ Güvenli bir çıktı yolu oluşturur (örn: dosya_conv.pdf) """
    base, _ = os.path.splitext(p)
    return base + "_conv" + new_ext

def txt_to_pdf(path, out_path):
    if not REPORTLAB_OK: raise RuntimeError("'reportlab' kütüphanesi sunucuda eksik!")
    c = canvas.Canvas(out_path, pagesize=A4); page_w, page_h = A4; margin = 50; y = page_h - margin; line_height = 14; max_width = page_w - (2 * margin)
    try: pdfmetrics.registerFont(TTFont('Verdana', 'verdana.ttf')); c.setFont("Verdana", 10)
    except: c.setFont("Helvetica", 10)
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        for line in f:
            original_line = line.rstrip("\n"); words = original_line.split(); line_to_draw = ""
            for word in words:
                if c.stringWidth(line_to_draw + word + " ", c._fontname, c._fontsize) <= max_width: line_to_draw += word + " "
                else:
                    if y < margin + line_height: c.showPage(); c.setFont(c._fontname, c._fontsize); y = page_h - margin
                    c.drawString(margin, y, line_to_draw.strip()); y -= line_height; line_to_draw = word + " "
            if line_to_draw.strip():
                if y < margin + line_height: c.showPage(); c.setFont(c._fontname, c._fontsize); y = page_h - margin
                c.drawString(margin, y, line_to_draw.strip()); y -= line_height
            if not original_line.strip(): y -= line_height
    c.save()

def pdf_to_txt(path, out_path):
    full_text = ""
    if PYPDF_OK:
        try:
            reader = PdfReader(path)
            for page in reader.pages: full_text += (page.extract_text() or "") + "\n"
            full_text = full_text.strip()
        except Exception: full_text = ""
    
    if len(full_text) < 100 and OCR_OK:
        print("PDF'ten metin okunamadı, OCR (Tesseract) deneniyor...")
        try:
            images = convert_from_path(path) # Poppler burada lazım
            ocr_text_list = []
            for img in images:
                ocr_text_list.append(pytesseract.image_to_string(img, lang='tur+eng')) # Tesseract burada lazım
            full_text = "\n".join(ocr_text_list).strip()
            print("OCR ile metin başarıyla çekildi.")
        except Exception as ocr_error:
            print(f"OCR hatası (Tesseract/Poppler kurulu mu?): {ocr_error}")
            raise RuntimeError(f"OCR hatası: Tesseract veya Poppler sunucuda kurulu değil! Hata: {ocr_error}")
    with open(out_path, "w", encoding="utf-8") as f: f.write(full_text)

def docx_to_txt(path, out_path):
    if not DOCX_OK: raise RuntimeError("'python-docx' kütüphanesi sunucuda eksik!")
    doc = Document(path)
    with open(out_path, "w", encoding="utf-8") as f:
        for p in doc.paragraphs: f.write(p.text + "\n")

def txt_to_docx(path, out_path):
    if not DOCX_OK: raise RuntimeError("'python-docx' kütüphanesi sunucuda eksik!")
    doc = Document()
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        for line in f: doc.add_paragraph(line.rstrip("\n"))
    doc.save(out_path)

def excel_to_docx(path, out_path):
    if not PANDAS_OK or not DOCX_OK: raise RuntimeError("'pandas', 'openpyxl' ve 'python-docx' kütüphaneleri eksik!")
    ext = os.path.splitext(path)[1].lower()
    if ext == ".csv":
        try: df = pd.read_csv(path, delimiter=';')
        except Exception: df = pd.read_csv(path)
    else: df = pd.read_excel(path)
    df = df.fillna(''); doc = Document()
    doc.add_heading(os.path.basename(path), level=1)
    table = doc.add_table(rows=1, cols=len(df.columns)); table.style = 'Table Grid'
    for i, col_name in enumerate(df.columns): table.cell(0, i).text = str(col_name)
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, cell_value in enumerate(row): row_cells[i].text = str(cell_value)
    doc.save(out_path)

def excel_to_txt(path, out_path):
    if not PANDAS_OK: raise RuntimeError("'pandas' ve 'openpyxl' kütüphaneleri eksik!")
    ext = os.path.splitext(path)[1].lower()
    if ext == ".csv":
        try: df = pd.read_csv(path, delimiter=';')
        except Exception: df = pd.read_csv(path)
    else: df = pd.read_excel(path)
    df.to_csv(out_path, sep='\t', index=False, header=True, encoding="utf-8-sig")

# --- YARDIMCILAR BİTTİ ---